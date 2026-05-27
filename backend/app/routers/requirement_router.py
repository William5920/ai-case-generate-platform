from fastapi import APIRouter, Depends, UploadFile, File, Query
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from typing import Optional
from urllib.parse import quote
from io import BytesIO

from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.schemas.requirement import (
    CreateRequirementRequest, UpdateRequirementRequest, RequirementListQuery
)
from app.services.requirement_service import requirement_service
from app.services.file_service import file_service
from app.services.split_service import split_service
from app.schemas.split import (
    ConfirmAndTestRequest, ExecuteSplitRequest, UpdateSplitRequest,
    AddSplitRequest
)

router = APIRouter(prefix="/requirements", tags=["需求管理"])


@router.post("", summary="创建需求")
async def create_requirement(
    req: CreateRequirementRequest,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await requirement_service.create_requirement(db, user_id, req)
    return {"code": 200, "message": "success", "data": data}


@router.get("", summary="获取需求列表")
async def get_requirement_list(
    pageNo: int = Query(1, ge=1),
    pageSize: int = Query(10, ge=1, le=100),
    keyword: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    query = RequirementListQuery(
        pageNo=pageNo, pageSize=pageSize,
        keyword=keyword, status=status
    )
    data = await requirement_service.get_requirement_list(db, user_id, query)
    return {"code": 200, "message": "success", "data": data}


@router.get("/{requirement_id}", summary="获取需求详情")
async def get_requirement_detail(
    requirement_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await requirement_service.get_requirement_detail(db, user_id, requirement_id)
    if not data:
        return {"code": 404, "message": "需求不存在", "data": None}
    return {"code": 200, "message": "success", "data": data}


@router.put("/{requirement_id}", summary="更新需求")
async def update_requirement(
    requirement_id: str,
    req: UpdateRequirementRequest,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await requirement_service.update_requirement(db, user_id, requirement_id, req)
    if not data:
        return {"code": 404, "message": "需求不存在", "data": None}
    return {"code": 200, "message": "success", "data": data}


@router.delete("/{requirement_id}", summary="删除需求")
async def delete_requirement(
    requirement_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    success = await requirement_service.delete_requirement(db, user_id, requirement_id)
    if not success:
        return {"code": 404, "message": "需求不存在", "data": None}
    return {"code": 200, "message": "success", "data": None}


@router.post("/upload", summary="上传需求文档")
async def upload_file(
    file: UploadFile = File(...),
    purpose: str = Query("requirement"),
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    try:
        data = await file_service.upload_file(db, user_id, file, purpose)
        return {"code": 200, "message": "success", "data": data}
    except ValueError as e:
        return {"code": 400, "message": str(e), "data": None}


@router.get("/{requirement_id}/export", summary="导出标准化文档")
async def export_document(
    requirement_id: str,
    format: str = Query("markdown", pattern=r"^(markdown|docx)$"),
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await requirement_service.get_requirement_detail(db, user_id, requirement_id)
    if not data:
        return {"code": 404, "message": "需求不存在", "data": None}

    title = data.get("title", "需求规格说明书")
    content = data.get("standardizedContent") or ""

    if format == "markdown":
        safe_filename = quote(f"{title}.md")
        return Response(
            content=content,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"}
        )

    if format == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument()
        lines = content.split("\n")
        for line in lines:
            trimmed = line.strip()
            if trimmed.startswith("# "):
                doc.add_heading(trimmed[2:], level=1)
            elif trimmed.startswith("## "):
                doc.add_heading(trimmed[3:], level=2)
            elif trimmed.startswith("### "):
                doc.add_heading(trimmed[4:], level=3)
            elif trimmed == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(trimmed)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        safe_filename = quote(f"{title}.docx")
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}"}
        )


@router.post("/{requirement_id}/split", summary="执行需求拆分")
async def execute_split(
    requirement_id: str,
    req: ExecuteSplitRequest,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    try:
        data = await split_service.execute_split(
            db, user_id,
            requirement_id=requirement_id,
            standardized_content=req.standardizedContent
        )
        return {"code": 200, "message": "success", "data": data}
    except ValueError as e:
        return {"code": 400, "message": str(e), "data": None}


@router.get("/{requirement_id}/splits", summary="获取拆分结果列表")
async def get_split_list(
    requirement_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await split_service.get_split_list(db, user_id, requirement_id)
    if not data:
        return {"code": 404, "message": "需求不存在", "data": None}
    return {"code": 200, "message": "success", "data": data}


@router.put("/{requirement_id}/splits/{split_id}", summary="更新拆分项")
async def update_split(
    requirement_id: str,
    split_id: str,
    req: UpdateSplitRequest,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await split_service.update_split(
        db, user_id, requirement_id, split_id,
        content=req.content, order=req.order
    )
    if not data:
        return {"code": 404, "message": "拆分项不存在", "data": None}
    return {"code": 200, "message": "success", "data": data}


@router.delete("/{requirement_id}/splits/{split_id}", summary="删除拆分项")
async def delete_split(
    requirement_id: str,
    split_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    success = await split_service.delete_split(db, user_id, requirement_id, split_id)
    if not success:
        return {"code": 404, "message": "拆分项不存在", "data": None}
    return {"code": 200, "message": "success", "data": None}


@router.post("/{requirement_id}/splits", summary="手动添加拆分项")
async def add_split(
    requirement_id: str,
    req: AddSplitRequest,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    data = await split_service.add_split(
        db, user_id, requirement_id,
        content=req.content, order=req.order
    )
    if not data:
        return {"code": 404, "message": "需求不存在", "data": None}
    return {"code": 200, "message": "success", "data": data}


@router.post("/{requirement_id}/confirm-and-test", summary="确认拆分并生成测试")
async def confirm_and_test(
    requirement_id: str,
    req: ConfirmAndTestRequest,
    db: AsyncSession = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user.id
    try:
        actual_id = None
        if requirement_id and requirement_id != "null":
            actual_id = requirement_id
        elif req.requirementId:
            actual_id = req.requirementId

        if not actual_id:
            from app.services.requirement_service import requirement_service
            from app.schemas.requirement import CreateRequirementRequest
            create_req = CreateRequirementRequest(
                title=req.title or "未命名需求",
                inputMode="text",
                rawContent=req.standardizedContent or "",
                templateId=req.templateId
            )
            create_data = await requirement_service.create_requirement(db, user_id, create_req)
            actual_id = create_data["id"]

        data = await split_service.confirm_and_test(
            db, user_id,
            requirement_id=actual_id,
            title=req.title or "未命名需求",
            split_requirements=req.splitRequirements or [],
            standardized_content=req.standardizedContent or "",
            template_id=req.templateId
        )
        return {"code": 200, "message": "success", "data": data}
    except ValueError as e:
        return {"code": 400, "message": str(e), "data": None}
